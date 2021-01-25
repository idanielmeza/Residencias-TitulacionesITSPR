from tkinter import *
from tkinter import ttk
import sqlite3
from tkinter import messagebox as ms
import shutil
import os
from tkinter import filedialog
from datetime import *
from docx import Document


class Tec:
    
    db_name = 'database.db'
    def __init__(self,ventana):
        self.wind= ventana
        self.wind.title('Bienvenido ITSPR')
        self.wind.iconbitmap('itspr.ico')
        self.wind.geometry('220x70')

        ttk.Button(self.wind, text='Maestro', command= self.login_maestro).pack(fill='both')
        Label(self.wind).pack()
        ttk.Button(self.wind, text='Alumno', command= self.login_alumno).pack(fill='both')

    def run_query(self,query, parametros = ()):
        with sqlite3.connect(self.db_name) as conn:
            cursor = conn.cursor()
            resultado = cursor.execute(query, parametros)
            conn.commit()
        return resultado
#### VETANAS MAESTROS
    def login_maestro(self):
        self.wind.destroy()
        self.loginm = Tk()
        self.loginm.title('Login Maestros')
        self.loginm.iconbitmap('itspr.ico')
        img = PhotoImage(file="itspr.gif").subsample(3,3)

        Label(self.loginm, image= img).pack()

        Label(self.loginm,text='Numero de control : ').pack(fill='both')
        self.user_m = Entry(self.loginm)
        self.user_m.pack()

        Label(self.loginm, text='Contraseña : ').pack(fill='both')

        self.cont_m = Entry(self.loginm, show = '*')
        self.cont_m.pack()

        Label(self.loginm).pack()

        ttk.Button(self.loginm, text='Conectar', command = self.conectar_m).pack(fill='both')

        self.loginm.mainloop()
        
    def conectar_m(self):
        
        u = self.user_m.get()
        c = self.cont_m.get()
        parametros = u,c
        query = 'select * from log where usuario = ? and pass = ?'
        datos = self.run_query(query,parametros)

        if datos.fetchall():
            self.inicio_maestro()
        else:
            ms.showerror('Conectado','Usuario Incorrecto')

    def inicio_maestro(self):
        
        self.loginm.destroy()
        self.wind2 = Tk()
        self.wind2.title('Registro Alumnos ITSPR')
        self.wind2.iconbitmap('itspr.ico')

        #NoteBook Maestro

        notebook = ttk.Notebook(self.wind2)
        notebook.pack(fill = 'both')
        ## NOTEBOOK PESTAÑA MEMORIAS
        self.memorias_wind = ttk.Frame(notebook)
        notebook.add(self.memorias_wind, text= 'Residencias')
### REGISTRO FRAME MEMORIAS       
        self.registro = LabelFrame(self.memorias_wind)
        self.registro.grid(row=0, column=0)

        ## FRAME REGISTRO 
        Label(self.registro, text = 'Agrega Alumnos').grid(row=0, column=0, columnspan=3)
        Label(self.registro).grid(row =1, column=2)

        Label(self.registro, text = 'Numero de Control : ').grid(row=2, column =0)
        self.num_mem = Entry(self.registro)
        self.num_mem.grid(row=2 , column=1)

        Label(self.registro, text='Contraseña : ').grid(row=3, column=0)
        self.cont_mem = Entry(self.registro)
        self.cont_mem.grid(row = 3, column = 1)

        Label(self.registro, text = 'Nombre : ').grid(row=4, column =0)
        self.nombre_mem = Entry(self.registro)
        self.nombre_mem.grid(row=4 , column=1)

        Label(self.registro, text = 'Apellido : ').grid(row=5, column =0)
        self.ap_mem = Entry(self.registro)
        self.ap_mem.grid(row=5 , column=1)

        Label(self.registro, text = 'Carrera : ').grid(row=6, column =0)
        self.carrera_mem = Entry(self.registro)
        self.carrera_mem.grid(row=6 , column=1)

        Label(self.registro, text = 'Maestro : ').grid(row=7, column =0)
        self.maestro_mem = Entry(self.registro)
        self.maestro_mem.grid(row=7 , column=1)

        #Label(self.registro, text = 'Archivo : ').grid(row=8, column =0)
        #self.archivo_mem = Entry(self.registro)
        #self.archivo_mem.grid(row=8 , column=1)
        #ttk.Button(self.registro, text = 'Seleccionar', command= self.seleccion_archivo).grid(row=8, column=2)

        ttk.Button(self.registro, text ='Agregar', command= self.agregar_memorias).grid(row = 9, column=1)
###########################################

        img = PhotoImage(file="itspr.gif").subsample(2,2)
        Label(self.memorias_wind, image= img).grid(row=0, column=1)

### BUSCAR FRAME MEMORIAS
        self.buscar = LabelFrame(self.memorias_wind)
        self.buscar.grid(row=0, column= 2)


        ##BOTON PRUEBA
        Label(self.buscar, text = 'Buscar Alumnos').grid(row=0, column=0, columnspan=3)
        Label(self.buscar).grid(row=1, column=0)

        Label(self.buscar, text='Numero de control : ').grid(row=2, column=0)
        self.buscar_num_mem = Entry(self.buscar)
        self.buscar_num_mem.grid(row=2, column = 1)
        ttk.Button(self.buscar, text='Buscar', command = self.buscar_m_n).grid(row=2, column=2)

        Label(self.buscar).grid(row=3,column=0)

        Label(self.buscar, text='Carrera : ').grid(row=4, column=0)
        self.buscar_carrera_mem = Entry(self.buscar)
        self.buscar_carrera_mem.grid(row=4, column = 1)
        ttk.Button(self.buscar, text='Buscar', command = self.buscar_m_c).grid(row=4, column=2)

        Label(self.buscar).grid(row=5, column = 0)

        Label(self.buscar, text = 'Apellido : ').grid(row=6, column=0)
        self.buscar_ap_mem = Entry(self.buscar)
        self.buscar_ap_mem.grid(row=6, column=1)
        ttk.Button(self.buscar, text='Buscar', command = self.buscar_m_a).grid(row=6, column=2)

        Label(self.buscar).grid(row=7,column=0)
        ttk.Button(self.buscar, text='Ver todos', command= self.obtener_memorias).grid(row=8, column=1)

### FRAME TREEVIEW MEMORIAS
        self.frame_tabla_mem = LabelFrame(self.memorias_wind)
        self.frame_tabla_mem.grid(row = 1, column=0, columnspan=3)

        self.tabla_memorias = ttk.Treeview(self.frame_tabla_mem, columns = ("#0","#1","#2","#3","#4"))
        self.tabla_memorias.grid(row = 0, column = 0, columnspan = 5)
        self.tabla_memorias.heading("#0", text= 'Num Control')
        self.tabla_memorias.heading("#1", text= 'Nombre')
        self.tabla_memorias.heading("#2", text= 'Apellido')
        self.tabla_memorias.heading("#3", text= 'Carrera')
        self.tabla_memorias.heading("#4", text= 'Maestro')
        self.tabla_memorias.heading("#5", text= 'Archivo')

        #BOTONES 
        ttk.Button(self.frame_tabla_mem, text='Editar', command= self.editar_memorias).grid(row =1, column=1)
        ttk.Button(self.frame_tabla_mem, text='Borrar', command= self.borrar_memoria).grid(row = 1, column=2)
        ttk.Button(self.frame_tabla_mem,text='Reporte', command = self.reporte_mem).grid(row = 1, column=3)

        ## NOTEBOOK TITULACIONES
        self.titulaciones_wind = ttk.Frame(notebook)
        notebook.add(self.titulaciones_wind, text= 'Titulaciones')

###FRAME TITULACIONES
        self.registro2 = LabelFrame(self.titulaciones_wind)
        self.registro2.grid(row=0, column=0)

        ## FRAME registro2 
        Label(self.registro2, text = 'Agrega Alumnos').grid(row=0, column=0, columnspan=3)
        Label(self.registro2).grid(row =1, column=2)

        Label(self.registro2, text = 'Numero de Control : ').grid(row=2, column =0)
        self.num_tit = Entry(self.registro2)
        self.num_tit.grid(row=2 , column=1)

        Label(self.registro2, text='Contraseña : ').grid(row=3, column=0)
        self.cont_tit = Entry(self.registro2)
        self.cont_tit.grid(row = 3, column = 1)

        Label(self.registro2, text = 'Nombre : ').grid(row=4, column =0)
        self.nombre_tit = Entry(self.registro2)
        self.nombre_tit.grid(row=4 , column=1)

        Label(self.registro2, text = 'Apellido : ').grid(row=5, column =0)
        self.ap_tit = Entry(self.registro2)
        self.ap_tit.grid(row=5 , column=1)

        Label(self.registro2, text = 'Carrera : ').grid(row=6, column =0)
        self.carrera_tit = Entry(self.registro2)
        self.carrera_tit.grid(row=6 , column=1)

        Label(self.registro2, text = 'Maestro : ').grid(row=7, column =0)
        self.maestro_tit = Entry(self.registro2)
        self.maestro_tit.grid(row=7 , column=1)

        #Label(self.registro2, text = 'Archivo : ').grid(row=8, column =0)
        #self.archivo_tit = Entry(self.registro2)
        #self.archivo_tit.grid(row=8 , column=1)
        #ttk.Button(self.registro2, text = 'Seleccionar').grid(row=8, column=2)

        ttk.Button(self.registro2, text ='Agregar', command= self.agregar_titulaciones).grid(row = 9, column=1)

        Label(self.titulaciones_wind, image= img).grid(row=0, column=1)

### BUSCAR FRAME titulaciones
        self.buscar = LabelFrame(self.titulaciones_wind)
        self.buscar.grid(row=0, column= 2)

        ##BOTON PRUEBA
        Label(self.buscar, text = 'Buscar Alumnos').grid(row=0, column=0, columnspan=3)
        Label(self.buscar).grid(row=1, column=0)

        Label(self.buscar, text='Numero de control : ').grid(row=2, column=0)
        self.buscar_num_tit = Entry(self.buscar)
        self.buscar_num_tit.grid(row=2, column = 1)
        ttk.Button(self.buscar, text='Buscar', command = self.buscar_t_n).grid(row=2, column=2)

        Label(self.buscar).grid(row =3, column=0)

        Label(self.buscar, text='Carrera : ').grid(row=4, column=0)
        self.buscar_carrera_tit = Entry(self.buscar)
        self.buscar_carrera_tit.grid(row=4, column = 1)
        ttk.Button(self.buscar, text='Buscar', command = self.buscar_t_c).grid(row=4, column=2)

        Label(self.buscar).grid(row=5, column = 0)

        Label(self.buscar, text = 'Apellido : ').grid(row=6, column=0)
        self.buscar_ap_tit = Entry(self.buscar)
        self.buscar_ap_tit.grid(row=6, column=1)
        ttk.Button(self.buscar, text='Buscar', command = self.buscar_t_a).grid(row=6, column=2)

        Label(self.buscar).grid(row=7,column=0)
        ttk.Button(self.buscar, text='Ver todos', command= self.obtener_titulaciones).grid(row=8, column=1)

### FRAME TREEVIEW titulaciones
        self.frame_tabla_tit = LabelFrame(self.titulaciones_wind)
        self.frame_tabla_tit.grid(row = 1, column=0, columnspan=3)

        self.tabla_titulaciones = ttk.Treeview(self.frame_tabla_tit, columns = ("#0","#1","#2","#3","#4"))
        self.tabla_titulaciones.grid(row = 0, column = 0, columnspan = 5)
        self.tabla_titulaciones.heading("#0", text= 'Num Control')
        self.tabla_titulaciones.heading("#1", text= 'Nombre')
        self.tabla_titulaciones.heading("#2", text= 'Apellido')
        self.tabla_titulaciones.heading("#3", text= 'Carrera')
        self.tabla_titulaciones.heading("#4", text= 'Maestro')
        self.tabla_titulaciones.heading("#5", text= 'Archivo')

        #BOTONES 
        ttk.Button(self.frame_tabla_tit, text='Editar', command = self.editar_titulaciones).grid(row =1, column=1)
        ttk.Button(self.frame_tabla_tit, text='Borrar', command = self.borrar_titulaciones).grid(row = 1, column=2)
        ttk.Button(self.frame_tabla_tit,text='Reporte', command = self.reporte_tit).grid(row = 1, column=3)        

        

        self.wind2.mainloop()

### REPORTES MEMORIAS
    def reporte_mem(self):
        self.rep_mem = Toplevel()
        self.rep_mem.title('Genera reportes')
        self.rep_mem.iconbitmap('itspr.ico')


        Label(self.rep_mem, text = 'Selecciona un Rango de fechas en formato AAAA-MM-DD', fg='red').grid(row = 0, column=0, columnspan=3)

        Label(self.rep_mem, text= 'Fecha de Inicio').grid(row=2, column = 0)
        Label(self.rep_mem, text= 'Fecha Final').grid(row =2, column = 1)
        Label(self.rep_mem, text = 'Carrera').grid(row = 2, column = 2)

        self.finicio_mem = Entry(self.rep_mem)
        self.finicio_mem.grid(row = 3, column=0)

        self.ffinal_mem = Entry(self.rep_mem)
        self.ffinal_mem.grid(row = 3, column=1)
        
        self.carrera_r_mem = Entry(self.rep_mem)
        self.carrera_r_mem.grid(row = 3, column=2)

        ttk.Button(self.rep_mem, text = 'Generar Reporte', command= self.reportes_memorias).grid(row = 4, column = 0, columnspan=3)

        # TREEVIEW REPORTES MEMORIAS

        self.tabla_reportes_mem = ttk.Treeview(self.rep_mem, columns=('#0','#1','#2'))
        self.tabla_reportes_mem.grid(row = 5, column=0, columnspan=3)
        self.tabla_reportes_mem.heading('#0', text='Numero de Control')
        self.tabla_reportes_mem.heading('#1', text='Titulo de Residencias')
        self.tabla_reportes_mem.heading('#2', text='Fecha de carga')
        self.tabla_reportes_mem.heading('#3', text='Carrera')

        # BOTON VER REPORTE

        #ttk.Button(self.rep_mem, text = 'Ver Reporte',command = self.ver_reporte).grid(row = 6, column=0, columnspan=3)

        self.rep_mem.mainloop()

### REPORTES TITULACIONES
    def reporte_tit(self):
        self.rep_tit = Toplevel()
        self.rep_tit.title('Genera reportes')
        self.rep_tit.iconbitmap('itspr.ico')


        Label(self.rep_tit, text = 'Selecciona un Rango de fechas en formato AAAA-MM-DD', fg='red').grid(row = 0, column=0, columnspan=3)

        Label(self.rep_tit, text= 'Fecha de Inicio').grid(row=2, column = 0)
        Label(self.rep_tit, text= 'Fecha Final').grid(row =2, column = 1)
        Label(self.rep_tit, text = 'Carrera').grid(row = 2, column = 2)

        self.finicio_tit = Entry(self.rep_tit)
        self.finicio_tit.grid(row = 3, column=0)

        self.ffinal_tit = Entry(self.rep_tit)
        self.ffinal_tit.grid(row = 3, column=1)
        
        self.carrera_r_tit = Entry(self.rep_tit)
        self.carrera_r_tit.grid(row = 3, column=2)

        ttk.Button(self.rep_tit, text = 'Generar Reporte', command= self.reportes_titulaciones).grid(row = 4, column = 0, columnspan=3)

        # TREEVIEW REPORTES titulaciones

        self.tabla_reportes_tit = ttk.Treeview(self.rep_tit, columns=('#0','#1','#2'))
        self.tabla_reportes_tit.grid(row = 5, column=0, columnspan=3)
        self.tabla_reportes_tit.heading('#0', text='Numero de Control')
        self.tabla_reportes_tit.heading('#1', text='Titulo de titulaciones')
        self.tabla_reportes_tit.heading('#2', text='Fecha de carga')
        self.tabla_reportes_tit.heading('#3', text='Carrera')

        # BOTON VER REPORTE

        #ttk.Button(self.rep_tit, text = 'Ver Reporte', command = self.ver_reporte).grid(row = 6, column=0, columnspan=3)

        self.rep_tit.mainloop()

## EDITAR ALUMNO MEMORIAS
    def editar_memorias(self):
        numcontrol = self.tabla_memorias.item(self.tabla_memorias.selection())['text']
        self.edit_mem = Toplevel()
        self.edit_mem.title('Editar alumno {}'.format(numcontrol))
        self.edit_mem.iconbitmap('itspr.ico')
        
        Label(self.edit_mem, text = 'Actualiza los datos del alumno {}'.format(numcontrol), font=('Open Sans',12)).grid(row=0, column=0, columnspan=2)

        Label(self.edit_mem, text ='Nombre : ').grid(row =1, column=0)
        self.edit_nombre_mem = Entry(self.edit_mem)
        self.edit_nombre_mem.grid(row= 1, column = 1)

        Label(self.edit_mem, text ='Apellido : ').grid(row =2, column=0)
        self.edit_apellido_mem = Entry(self.edit_mem)
        self.edit_apellido_mem.grid(row= 2, column = 1)

        Label(self.edit_mem, text ='Contraseña : ').grid(row =3, column=0)
        self.edit_cont_mem = Entry(self.edit_mem)
        self.edit_cont_mem.grid(row= 3, column = 1)

        Label(self.edit_mem, text ='Carrera : ').grid(row =4, column=0)
        self.edit_carrera_mem = Entry(self.edit_mem)
        self.edit_carrera_mem.grid(row= 4, column = 1)

        Label(self.edit_mem, text ='Maestro : ').grid(row =5, column=0)
        self.edit_maestro_mem = Entry(self.edit_mem)
        self.edit_maestro_mem.grid(row= 5, column = 1)

        ttk.Button(self.edit_mem, text='Actualizar', command = lambda: self.update_mem(numcontrol)).grid(row =6, column =1)

## EDITAR ALUMNO titulaciones
    def editar_titulaciones(self):
        numcontrol = self.tabla_titulaciones.item(self.tabla_titulaciones.selection())['text']
        self.edit_tit = Toplevel()
        self.edit_tit.title('Editar alumno {}'.format(numcontrol))
        self.edit_tit.iconbitmap('itspr.ico')
        
        Label(self.edit_tit, text = 'Actualiza los datos del alumno {}'.format(numcontrol), font=('Open Sans',12)).grid(row=0, column=0, columnspan=2)

        Label(self.edit_tit, text ='Nombre : ').grid(row =1, column=0)
        self.edit_nombre_tit = Entry(self.edit_tit)
        self.edit_nombre_tit.grid(row= 1, column = 1)

        Label(self.edit_tit, text ='Apellido : ').grid(row =2, column=0)
        self.edit_apellido_tit = Entry(self.edit_tit)
        self.edit_apellido_tit.grid(row= 2, column = 1)

        Label(self.edit_tit, text ='Contraseña : ').grid(row =3, column=0)
        self.edit_cont_tit = Entry(self.edit_tit)
        self.edit_cont_tit.grid(row= 3, column = 1)

        Label(self.edit_tit, text ='Carrera : ').grid(row =4, column=0)
        self.edit_carrera_tit = Entry(self.edit_tit)
        self.edit_carrera_tit.grid(row= 4, column = 1)

        Label(self.edit_tit, text ='Maestro : ').grid(row =5, column=0)
        self.edit_maestro_tit = Entry(self.edit_tit)
        self.edit_maestro_tit.grid(row= 5, column = 1)

        ttk.Button(self.edit_tit, text='Actualizar', command = lambda: self.update_tit(numcontrol)).grid(row =6, column =1)

#### VENTANAS ALUMNOS
    def login_alumno(self):
        self.wind.destroy()
        self.logina = Tk()
        self.logina.title('Login Alumnos')
        self.logina.iconbitmap('itspr.ico')
        #self.logina.geometry('250x180')

        notebook2 = ttk.Notebook(self.logina)
        notebook2.pack(fill = 'both')
        
        self.logina_notebook = LabelFrame(notebook2)
        notebook2.add(self.logina_notebook,text='Residencias')

        self.logina2_notebook = LabelFrame(notebook2)

        notebook2.add(self.logina2_notebook, text ='Tiutlaciones')
        img = PhotoImage(file= "itspr.gif").subsample(3,3)
        ##Memorias Login Alumnos
        
        Label(self.logina_notebook, image= img).pack()

        Label(self.logina_notebook, text ='Numero de Control : ').pack(fill='both')
        self.usuario_mem_alumno = Entry(self.logina_notebook)
        self.usuario_mem_alumno.pack()

        Label(self.logina_notebook, text ='Contraseña : ').pack(fill='both')
        self.cont_mem_alumno = Entry(self.logina_notebook, show='*')
        self.cont_mem_alumno.pack()

        Label(self.logina_notebook).pack()
        ttk.Button(self.logina_notebook, text='Conectar', command= self.conectar_alumnos_memorias).pack(fill='both')
        ttk.Button(self.logina_notebook, text = 'Olvide mi contraseña', command = self.reset_password).pack(fill='both')

        ##Titulaciones Login Alumnos
        Label(self.logina2_notebook, image= img).pack()

        Label(self.logina2_notebook, text ='Numero de Control : ').pack(fill='both')
        self.usuario_tit_alumno = Entry(self.logina2_notebook)
        self.usuario_tit_alumno.pack()

        Label(self.logina2_notebook, text ='Contraseña : ').pack(fill='both')
        self.cont_tit_alumno = Entry(self.logina2_notebook, show='*')
        self.cont_tit_alumno.pack()

        Label(self.logina2_notebook).pack()
        ttk.Button(self.logina2_notebook, text='Conectar', command = self.conectar_alumnos_titulaciones).pack(fill='both')
        ttk.Button(self.logina2_notebook, text = 'Olvide mi contraseña', command = self.reset_password).pack(fill='both')

        self.logina.mainloop()
    ## INICIO ALUMNOS

    def inicio_memorias(self):
        
        usuario = self.usuario_mem_alumno.get()
        self.logina.destroy()
        self.inicio_mem_alumnos = Tk()
        self.inicio_mem_alumnos.title('Residencias ITSPR')
        self.inicio_mem_alumnos.iconbitmap('itspr.ico')

        
        query = 'select nombre, apellido, carrera, maestro from memorias where numcontrol = ?'
        datos = self.run_query(query,(usuario,))
        query2 = 'select nombre, apellido from memorias where numcontrol = ?'
        datos2 = self.run_query(query2,(usuario,))
        
        nom1= ''
        ap = ''
        for (nombre, apellido) in datos2:
            nom1 = nombre
            ap = apellido
        
        Label(self.inicio_mem_alumnos, text = 'Bienvenido {} {}'.format(nom1,ap),font = ('Open Sans',12)).grid(row=0, column=0, columnspan=3)

        Label(self.inicio_mem_alumnos, text='Numero de control : ').grid(row = 1, column = 0)
        self.numcontrol_inicio_mem =Entry(self.inicio_mem_alumnos, state= NORMAL)
        self.numcontrol_inicio_mem.grid(row=1, column=1)
        self.numcontrol_inicio_mem.insert(0,usuario)

        Label(self.inicio_mem_alumnos, text='Nombre : ').grid(row = 2, column = 0)
        self.nombre_inicio_mem =Entry(self.inicio_mem_alumnos, state= NORMAL)
        self.nombre_inicio_mem.grid(row=2, column=1)
        
        Label(self.inicio_mem_alumnos, text='Apellido : ').grid(row = 3, column = 0)
        self.apellido_inicio_mem =Entry(self.inicio_mem_alumnos, state= NORMAL)
        self.apellido_inicio_mem.grid(row=3, column=1)

        Label(self.inicio_mem_alumnos, text='Carrera : ').grid(row = 4, column = 0)
        self.carrera_inicio_mem =Entry(self.inicio_mem_alumnos, state= NORMAL)
        self.carrera_inicio_mem.grid(row=4, column=1)

        Label(self.inicio_mem_alumnos, text='Maestro : ').grid(row = 5, column = 0)
        self.maestro_inicio_mem =Entry(self.inicio_mem_alumnos, state= NORMAL)
        self.maestro_inicio_mem.grid(row=5, column=1)

        Label(self.inicio_mem_alumnos, text='Titutlo del Documento : ').grid(row = 6, column = 0)
        self.nomarchivo_inicio_mem =Entry(self.inicio_mem_alumnos)
        self.nomarchivo_inicio_mem.grid(row=6, column=1)

        Label(self.inicio_mem_alumnos, text='Archivo : ').grid(row = 7, column = 0)
        self.archivo_inicio_mem =Entry(self.inicio_mem_alumnos)
        self.archivo_inicio_mem.grid(row=7, column=1)
        ttk.Button(self.inicio_mem_alumnos, text='Seleccionar', command= self.seleccion_archivo_mem).grid(row=7, column=2)

        ttk.Button(self.inicio_mem_alumnos, text='Cargar', command = self.cargar_mem).grid(row=8, column=1)

        for (nombre, apellido, carrera, maestro) in datos:
            self.nombre_inicio_mem.insert(0,nombre)
            self.apellido_inicio_mem.insert(0,apellido)
            self.carrera_inicio_mem.insert(0,carrera)
            self.maestro_inicio_mem.insert(0,maestro)

            self.numcontrol_inicio_mem.config(state = DISABLED)
            self.nombre_inicio_mem.config(state = DISABLED)
            self.apellido_inicio_mem.config(state = DISABLED)
            self.carrera_inicio_mem.config(state = DISABLED)
            self.maestro_inicio_mem.config(state = DISABLED)
            

        self.inicio_mem_alumnos.mainloop()

    def inicio_titulaciones(self):
        
        usuario = self.usuario_tit_alumno.get()
        self.logina.destroy()
        self.inicio_tit_alumnos = Tk()
        self.inicio_tit_alumnos.title('Titulaciones ITSPR')
        self.inicio_tit_alumnos.iconbitmap('itspr.ico')

        
        query = 'select nombre, apellido, carrera, maestro from titulaciones where numcontrol = ?'
        datos = self.run_query(query,(usuario,))
        query2 = 'select nombre, apellido from titulaciones where numcontrol = ?'
        datos2 = self.run_query(query2,(usuario,))
        
        nom1= ''
        ap = ''
        for (nombre, apellido) in datos2:
            nom1 = nombre
            ap = apellido
        
        Label(self.inicio_tit_alumnos, text = 'Bienvenido {} {}'.format(nom1,ap),font = ('Open Sans',12)).grid(row=0, column=0, columnspan=3)

        Label(self.inicio_tit_alumnos, text='Numero de control : ').grid(row = 1, column = 0)
        self.numcontrol_inicio_tit =Entry(self.inicio_tit_alumnos, state= NORMAL)
        self.numcontrol_inicio_tit.grid(row=1, column=1)
        self.numcontrol_inicio_tit.insert(0,usuario)

        Label(self.inicio_tit_alumnos, text='Nombre : ').grid(row = 2, column = 0)
        self.nombre_inicio_tit =Entry(self.inicio_tit_alumnos, state= NORMAL)
        self.nombre_inicio_tit.grid(row=2, column=1)
        
        Label(self.inicio_tit_alumnos, text='Apellido : ').grid(row = 3, column = 0)
        self.apellido_inicio_tit =Entry(self.inicio_tit_alumnos, state= NORMAL)
        self.apellido_inicio_tit.grid(row=3, column=1)

        Label(self.inicio_tit_alumnos, text='Carrera : ').grid(row = 4, column = 0)
        self.carrera_inicio_tit =Entry(self.inicio_tit_alumnos, state= NORMAL)
        self.carrera_inicio_tit.grid(row=4, column=1)

        Label(self.inicio_tit_alumnos, text='Maestro : ').grid(row = 5, column = 0)
        self.maestro_inicio_tit =Entry(self.inicio_tit_alumnos, state= NORMAL)
        self.maestro_inicio_tit.grid(row=5, column=1)

        Label(self.inicio_tit_alumnos, text='Titutlo del Documento : ').grid(row = 6, column = 0)
        self.nomarchivo_inicio_tit =Entry(self.inicio_tit_alumnos)
        self.nomarchivo_inicio_tit.grid(row=6, column=1)

        Label(self.inicio_tit_alumnos, text='Archivo : ').grid(row = 7, column = 0)
        self.archivo_inicio_tit =Entry(self.inicio_tit_alumnos)
        self.archivo_inicio_tit.grid(row=7, column=1)
        ttk.Button(self.inicio_tit_alumnos, text='Seleccionar', command= self.seleccion_archivo_tit).grid(row=7, column=2)

        ttk.Button(self.inicio_tit_alumnos, text='Cargar', command = self.cargar_tit).grid(row=8, column=1)

        for (nombre, apellido, carrera, maestro) in datos:
            self.nombre_inicio_tit.insert(0,nombre)
            self.apellido_inicio_tit.insert(0,apellido)
            self.carrera_inicio_tit.insert(0,carrera)
            self.maestro_inicio_tit.insert(0,maestro)

            self.numcontrol_inicio_tit.config(state = DISABLED)
            self.nombre_inicio_tit.config(state = DISABLED)
            self.apellido_inicio_tit.config(state = DISABLED)
            self.carrera_inicio_tit.config(state = DISABLED)
            self.maestro_inicio_tit.config(state = DISABLED)
            

        self.inicio_tit_alumnos.mainloop()

#### FUNCIONES MAESTROS
            
    def obtener_memorias(self):
        #Limpiando la tabla antes de la consulta
        registro = self.tabla_memorias.get_children()
        for elemento in registro:
            self.tabla_memorias.delete(elemento)
        query = 'select `numcontrol`, `nombre`, `apellido`, `carrera`,`maestro`, `archivo` from memorias'
        datos = self.run_query(query)
        for (numcontrol,nombre,apellido, carrera,maestro, archivo) in datos:
            self.tabla_memorias.insert('', 0, text = numcontrol, values = (nombre,apellido,carrera,maestro,archivo))    

    def obtener_titulaciones(self):
        #Limpiando la tabla antes de la consulta
        registro = self.tabla_titulaciones.get_children()
        for elemento in registro:
            self.tabla_titulaciones.delete(elemento)
        query = 'select `numcontrol`, `nombre`, `apellido`, `carrera`,`maestro`, `archivo` from titulaciones'
        datos = self.run_query(query)
        for (numcontrol,nombre,apellido, carrera,maestro, archivo) in datos:
            self.tabla_titulaciones.insert('', 0, text = numcontrol, values = (nombre,apellido,carrera,maestro,archivo))  

    def agregar_memorias(self):
        num = self.num_mem.get()
        cont = self.cont_mem.get()
        nombre = self.nombre_mem.get()
        apellido = self.ap_mem.get()
        carrera = self.carrera_mem.get()
        maestro = self.maestro_mem.get()
        #archivo = self.archivo_mem.get()

        
        
        #os.makedirs('C:\\ITSPR\\Memorias\\{}'.format(num), exist_ok=True)
        #archivo_ub = 'C:\\ITSPR\\Memorias\\{}'.format(num)
        #shutil.copy(archivo, archivo_ub)        

        

        parametros = num, cont, nombre, apellido, carrera, maestro
        query = 'insert into memorias(numcontrol, pass, nombre, apellido, carrera, maestro) values(?,?,?,?,?,?)'

        self.run_query(query,parametros)
        ms.showinfo('Alumno agregrado','El alumno {} {} ha sido agregado correctamente'.format(nombre,apellido))
        self.obtener_memorias()
        self.limpiar()
    
    def agregar_titulaciones(self):
        num = self.num_tit.get()
        cont = self.cont_tit.get()
        nombre = self.nombre_tit.get()
        apellido = self.ap_tit.get()
        carrera = self.carrera_tit.get()
        maestro = self.maestro_tit.get()    
        
        parametros = num, cont, nombre, apellido, carrera, maestro
        query = 'insert into titulaciones(numcontrol, pass, nombre, apellido, carrera, maestro) values(?,?,?,?,?,?)'

        self.run_query(query,parametros)
        ms.showinfo('Alumno agregrado','El alumno {} {} ha sido agregado correctamente'.format(nombre,apellido))
        self.obtener_titulaciones()
        self.limpiar()

    def borrar_memoria(self):
        bor = self.tabla_memorias.item(self.tabla_memorias.selection())['text']
        query = 'delete from memorias where numcontrol = ?'
        self.run_query(query,(bor,))
        self.obtener_memorias()

    def borrar_titulaciones(self):
        
        bor = self.tabla_titulaciones.item(self.tabla_titulaciones.selection())['text']

        query = 'delete from titulaciones where numcontrol = ?'
        self.run_query(query,(bor,))
        self.obtener_titulaciones()

    def buscar_m_n(self):
        #Limpiando la tabla antes de la consulta
        registro = self.tabla_memorias.get_children()
        for elemento in registro:
            self.tabla_memorias.delete(elemento)
        query = 'select `numcontrol`, `nombre`, `apellido`, `carrera`,`maestro`, `archivo` from memorias where numcontrol = ?'
        bor=self.buscar_num_mem.get()
        datos = self.run_query(query,(bor,))
        for (numcontrol,nombre,apellido, carrera,maestro, archivo) in datos:
            self.tabla_memorias.insert('', 0, text = numcontrol, values = (nombre,apellido,carrera,maestro,archivo))

    def buscar_m_c(self):
        #Limpiando la tabla antes de la consulta
        registro = self.tabla_memorias.get_children()
        for elemento in registro:
            self.tabla_memorias.delete(elemento)
        query = 'select `numcontrol`, `nombre`, `apellido`, `carrera`,`maestro`, `archivo` from memorias where carrera like ?'
        bor=self.buscar_carrera_mem.get() + '%'
        datos = self.run_query(query,(bor,))
        for (numcontrol,nombre,apellido, carrera,maestro, archivo) in datos:
            self.tabla_memorias.insert('', 0, text = numcontrol, values = (nombre,apellido,carrera,maestro,archivo))

    def buscar_m_a(self):
        #Limpiando la tabla antes de la consulta
        registro = self.tabla_memorias.get_children()
        for elemento in registro:
            self.tabla_memorias.delete(elemento)
        query = 'select `numcontrol`, `nombre`, `apellido`, `carrera`,`maestro`, `archivo` from memorias where apellido like ?'
        bor='%' + self.buscar_ap_mem.get() + '%'
        datos = self.run_query(query,(bor,))
        for (numcontrol,nombre,apellido, carrera,maestro, archivo) in datos:
            self.tabla_memorias.insert('', 0, text = numcontrol, values = (nombre,apellido,carrera,maestro,archivo))

    def buscar_t_n(self):
        #Limpiando la tabla antes de la consulta
        registro = self.tabla_titulaciones.get_children()
        for elemento in registro:
            self.tabla_titulaciones.delete(elemento)
        query = 'select `numcontrol`, `nombre`, `apellido`, `carrera`,`maestro`, `archivo` from titulaciones where numcontrol = ?'
        bor=self.buscar_num_tit.get()
        datos = self.run_query(query,(bor,))
        for (numcontrol,nombre,apellido, carrera,maestro, archivo) in datos:
            self.tabla_titulaciones.insert('', 0, text = numcontrol, values = (nombre,apellido,carrera,maestro,archivo))

    def buscar_t_c(self):
        #Limpiando la tabla antes de la consulta
        registro = self.tabla_titulaciones.get_children()
        for elemento in registro:
            self.tabla_titulaciones.delete(elemento)
        query = 'select `numcontrol`, `nombre`, `apellido`, `carrera`,`maestro`, `archivo` from titulaciones where carrera like ?'
        bor=self.buscar_carrera_tit.get() + '%'
        datos = self.run_query(query,(bor,))
        for (numcontrol,nombre,apellido, carrera,maestro, archivo) in datos:
            self.tabla_titulaciones.insert('', 0, text = numcontrol, values = (nombre,apellido,carrera,maestro,archivo))

    def buscar_t_a(self):
        #Limpiando la tabla antes de la consulta
        registro = self.tabla_titulaciones.get_children()
        for elemento in registro:
            self.tabla_titulaciones.delete(elemento)
        query = 'select `numcontrol`, `nombre`, `apellido`, `carrera`,`maestro`, `archivo` from titulaciones where apellido like ?'
        bor=self.buscar_ap_tit.get() + '%'
        datos = self.run_query(query,(bor,))
        for (numcontrol,nombre,apellido, carrera,maestro, archivo) in datos:
            self.tabla_titulaciones.insert('', 0, text = numcontrol, values = (nombre,apellido,carrera,maestro,archivo))

    def limpiar(self):
            self.num_mem.delete(0, END)
            self.cont_mem.delete(0, END)
            self.nombre_mem.delete(0, END)
            self.ap_mem.delete(0, END)
            self.carrera_mem.delete(0, END)
            self.maestro_mem.delete(0, END)
            #self.archivo_mem.delete(0,END)
            self.num_tit.delete(0, END)
            self.cont_tit.delete(0, END)
            self.nombre_tit.delete(0, END)
            self.ap_tit.delete(0, END)
            self.carrera_tit.delete(0, END)
            self.maestro_tit.delete(0, END)
            #self.archivo_tit.delete(0,END)

    def reportes_memorias(self):
        fechaInicio = self.finicio_mem.get()
        fechaFinal = self.ffinal_mem.get()
        carrera = self.carrera_r_mem.get()
        fechaActual = date.today()

        if fechaInicio and fechaFinal != '' and carrera == '':
            parametros = self.finicio_mem.get(), self.ffinal_mem.get()
            query = "select numcontrol,nomarchivo,fecha,carrera from memorias where fecha BETWEEN ? AND ? order by fecha desc"
            datos = self.run_query(query,parametros)

            #Limpiando la tabla antes de la consulta
            registro = self.tabla_reportes_mem.get_children()
            for elemento in registro:
                self.tabla_reportes_mem.delete(elemento)
            for(numcontrol, nomarchivo, fecha, carrera) in datos:
                self.tabla_reportes_mem.insert('', 0, text = numcontrol , values= (nomarchivo,fecha,carrera))

            query2 = "select numcontrol,nombre,apellido,carrera,maestro,nomarchivo,fecha from memorias where fecha BETWEEN ? AND ? order by fecha desc"
            datos2 = self.run_query(query2,parametros)

            document = Document()
            document.add_heading('Reporte memorias {} fecha : {} a {}'.format(fechaActual, fechaInicio, fechaFinal), 0)
            table = document.add_table(rows=1, cols =7)
            table.rows[0].cells[0].text = 'Numero de control'
            table.rows[0].cells[1].text = 'Nombre'
            table.rows[0].cells[2].text = 'Apellido'
            table.rows[0].cells[3].text = 'Carrera'
            table.rows[0].cells[4].text = 'Maestro'
            table.rows[0].cells[5].text = 'Titulo de Memorias'
            table.rows[0].cells[6].text = 'Fecha de carga'

            for (numcontrol,nombre,apellido,carrera,maestro,nomarchivo,fecha) in datos2:
                row_cells = table.add_row().cells
                row_cells[0].text = numcontrol
                row_cells[1].text = nombre
                row_cells[2].text = apellido
                row_cells[3].text = carrera
                row_cells[4].text = maestro
                row_cells[5].text = nomarchivo
                row_cells[6].text = fecha

            document.save('REPORTE MEMORIAS {} RANGO {} A {}.docx'.format(fechaActual,fechaInicio,fechaFinal))
            path = os.path.realpath('REPORTE MEMORIAS {} RANGO {} A {}.docx'.format(fechaActual,fechaInicio,fechaFinal))
            os.startfile(path)

        elif carrera != '' and (fechaInicio and fechaFinal) == '':
            query = 'select numcontrol,nomarchivo,fecha,carrera from memorias where fecha is not NULL and carrera like ?'
            datos = self.run_query(query,(self.carrera_r_mem.get(),))

                #Limpiando la tabla antes de la consulta
            registro = self.tabla_reportes_mem.get_children()
            for elemento in registro:
                self.tabla_reportes_mem.delete(elemento)
            for(numcontrol, nomarchivo, fecha, carrera) in datos:
                self.tabla_reportes_mem.insert('', 0, text = numcontrol , values= (nomarchivo,fecha,carrera))

            query2 = 'select numcontrol,nombre,apellido,carrera,maestro,nomarchivo,fecha from memorias where fecha is not NULL and carrera like ?'
            datos2 = self.run_query(query2,(self.carrera_r_mem.get(),))

            document = Document()
            document.add_heading('Reporte memorias {} de la carrera {}'.format(fechaActual, carrera), 0)
            table = document.add_table(rows=1, cols =7)
            table.rows[0].cells[0].text = 'Numero de control'
            table.rows[0].cells[1].text = 'Nombre'
            table.rows[0].cells[2].text = 'Apellido'
            table.rows[0].cells[3].text = 'Carrera'
            table.rows[0].cells[4].text = 'Maestro'
            table.rows[0].cells[5].text = 'Titulo de Memorias'
            table.rows[0].cells[6].text = 'Fecha de carga'

            for (numcontrol,nombre,apellido,carrera,maestro,nomarchivo,fecha) in datos2:
                row_cells = table.add_row().cells
                row_cells[0].text = numcontrol
                row_cells[1].text = nombre
                row_cells[2].text = apellido
                row_cells[3].text = carrera
                row_cells[4].text = maestro
                row_cells[5].text = nomarchivo
                row_cells[6].text = fecha

            document.save('REPORTE MEMORIAS {} CARRERA {}.docx'.format(fechaActual,carrera))
            path = os.path.realpath('REPORTE MEMORIAS {} CARRERA {}.docx'.format(fechaActual,carrera))
            os.startfile(path)
        
        elif fechaInicio and fechaFinal != '' and carrera != '':
            parametros = self.finicio_mem.get(), self.ffinal_mem.get(), self.carrera_r_mem.get()
            query = "select numcontrol,nomarchivo,fecha,carrera from memorias where fecha BETWEEN ? AND ? AND carrera like ?"
            datos = self.run_query(query, parametros)

            #Limpiando la tabla antes de la consulta
            registro = self.tabla_reportes_mem.get_children()
            for elemento in registro:
                self.tabla_reportes_mem.delete(elemento)
            for(numcontrol, nomarchivo, fecha, carrera) in datos:
                self.tabla_reportes_mem.insert('', 0, text = numcontrol , values= (nomarchivo,fecha,carrera))

            query2 = "select numcontrol,nombre,apellido,carrera,maestro,nomarchivo,fecha from memorias where fecha BETWEEN ? AND ? AND carrera like ?"
            datos2 = self.run_query(query2,parametros)

            document = Document()
            document.add_heading('Reporte memorias {} fecha : {} a {} de la carrera {}'.format(fechaActual, fechaInicio, fechaFinal,carrera), 0)
            table = document.add_table(rows=1, cols =7)
            table.rows[0].cells[0].text = 'Numero de control'
            table.rows[0].cells[1].text = 'Nombre'
            table.rows[0].cells[2].text = 'Apellido'
            table.rows[0].cells[3].text = 'Carrera'
            table.rows[0].cells[4].text = 'Maestro'
            table.rows[0].cells[5].text = 'Titulo de Memorias'
            table.rows[0].cells[6].text = 'Fecha de carga'

            for (numcontrol,nombre,apellido,carrera,maestro,nomarchivo,fecha) in datos2:
                row_cells = table.add_row().cells
                row_cells[0].text = numcontrol
                row_cells[1].text = nombre
                row_cells[2].text = apellido
                row_cells[3].text = carrera
                row_cells[4].text = maestro
                row_cells[5].text = nomarchivo
                row_cells[6].text = fecha

            document.save('REPORTE MEMORIAS {} CARRERA {} RANGO {} A {}.docx'.format(fechaActual,carrera,fechaInicio,fechaFinal))
            path = os.path.realpath('REPORTE MEMORIAS {} CARRERA {} RANGO {} A {}.docx'.format(fechaActual,carrera,fechaInicio,fechaFinal))
            os.startfile(path)

        else:
            ms.showerror('ERROR','Ingresa un rango de fechas o la carrera')

    def reportes_titulaciones(self):
        fechaInicio = self.finicio_tit.get()
        fechaFinal = self.ffinal_tit.get()
        carrera = self.carrera_r_tit.get()
        fechaActual = date.today()

        if fechaInicio and fechaFinal != '' and carrera == '':
            parametros = self.finicio_tit.get(), self.ffinal_tit.get()
            query = "select numcontrol,nomarchivo,fecha,carrera from titulaciones where fecha BETWEEN ? AND ? order by fecha desc"
            datos = self.run_query(query,parametros)

                #Limpiando la tabla antes de la consulta
            registro = self.tabla_reportes_tit.get_children()
            for elemento in registro:
                self.tabla_reportes_tit.delete(elemento)
            for(numcontrol, nomarchivo, fecha, carrera) in datos:
                self.tabla_reportes_tit.insert('', 0, text = numcontrol , values= (nomarchivo,fecha,carrera))
            
            query2 = "select numcontrol,nombre,apellido,carrera,maestro,nomarchivo,fecha from titulaciones where fecha BETWEEN ? AND ? order by fecha desc"
            datos2 = self.run_query(query2,parametros)

            document = Document()
            document.add_heading('Reporte titulaciones {} fecha : {} a {}'.format(fechaActual, fechaInicio, fechaFinal), 0)
            table = document.add_table(rows=1, cols =7)
            table.rows[0].cells[0].text = 'Numero de control'
            table.rows[0].cells[1].text = 'Nombre'
            table.rows[0].cells[2].text = 'Apellido'
            table.rows[0].cells[3].text = 'Carrera'
            table.rows[0].cells[4].text = 'Maestro'
            table.rows[0].cells[5].text = 'Titulo de titulaciones'
            table.rows[0].cells[6].text = 'Fecha de carga'

            for (numcontrol,nombre,apellido,carrera,maestro,nomarchivo,fecha) in datos2:
                row_cells = table.add_row().cells
                row_cells[0].text = numcontrol
                row_cells[1].text = nombre
                row_cells[2].text = apellido
                row_cells[3].text = carrera
                row_cells[4].text = maestro
                row_cells[5].text = nomarchivo
                row_cells[6].text = fecha

            document.save('REPORTE TITULACIONES {} RANGO {} A {}.docx'.format(fechaActual,fechaInicio,fechaFinal))
            path = os.path.realpath('REPORTE TITULACIONES {} RANGO {} A {}.docx'.format(fechaActual,fechaInicio,fechaFinal))
            os.startfile(path)

        elif carrera != '' and (fechaInicio and fechaFinal) == '':
            query = 'select numcontrol,nomarchivo,fecha,carrera from titulaciones where fecha is not NULL and carrera like ?'
            datos = self.run_query(query,(self.carrera_r_tit.get(),))

                #Limpiando la tabla antes de la consulta
            registro = self.tabla_reportes_tit.get_children()
            for elemento in registro:
                self.tabla_reportes_tit.delete(elemento)
            for(numcontrol, nomarchivo, fecha, carrera) in datos:
                self.tabla_reportes_tit.insert('', 0, text = numcontrol , values= (nomarchivo,fecha,carrera))
            
            query2 = 'select numcontrol,nombre,apellido,carrera,maestro,nomarchivo,fecha from titulaciones where fecha is not NULL and carrera like ?'
            datos2 = self.run_query(query2,(self.carrera_r_tit.get(),))

            document = Document()
            document.add_heading('Reporte titulaciones {} de la carrera {}'.format(fechaActual, carrera), 0)
            table = document.add_table(rows=1, cols =7)
            table.rows[0].cells[0].text = 'Numero de control'
            table.rows[0].cells[1].text = 'Nombre'
            table.rows[0].cells[2].text = 'Apellido'
            table.rows[0].cells[3].text = 'Carrera'
            table.rows[0].cells[4].text = 'Maestro'
            table.rows[0].cells[5].text = 'Titulo de titulaciones'
            table.rows[0].cells[6].text = 'Fecha de carga'

            for (numcontrol,nombre,apellido,carrera,maestro,nomarchivo,fecha) in datos2:
                row_cells = table.add_row().cells
                row_cells[0].text = numcontrol
                row_cells[1].text = nombre
                row_cells[2].text = apellido
                row_cells[3].text = carrera
                row_cells[4].text = maestro
                row_cells[5].text = nomarchivo
                row_cells[6].text = fecha

            document.save('REPORTE TITULACIONES {} CARRERA {}.docx'.format(fechaActual,carrera))
            path = os.path.realpath('REPORTE TITULACIONES {} CARRERA {}.docx'.format(fechaActual,carrera))
            os.startfile(path)
        
        elif fechaInicio and fechaFinal != '' and carrera != '':
            parametros = self.finicio_tit.get(), self.ffinal_tit.get(), self.carrera_r_tit.get()
            query = "select numcontrol,nomarchivo,fecha,carrera from titulaciones where fecha BETWEEN ? AND ? AND carrera like ?"
            datos = self.run_query(query, parametros)

            #Limpiando la tabla antes de la consulta
            registro = self.tabla_reportes_tit.get_children()
            for elemento in registro:
                self.tabla_reportes_tit.delete(elemento)
            for(numcontrol, nomarchivo, fecha, carrera) in datos:
                self.tabla_reportes_tit.insert('', 0, text = numcontrol , values= (nomarchivo,fecha,carrera))
            
            query2 = "select numcontrol,nombre,apellido,carrera,maestro,nomarchivo,fecha from titulaciones where fecha BETWEEN ? AND ? AND carrera like ?"
            datos2 = self.run_query(query2,parametros)

            document = Document()
            document.add_heading('Reporte titulaciones {} fecha : {} a {} de la carrera {}'.format(fechaActual, fechaInicio, fechaFinal,carrera), 0)
            table = document.add_table(rows=1, cols =7)
            table.rows[0].cells[0].text = 'Numero de control'
            table.rows[0].cells[1].text = 'Nombre'
            table.rows[0].cells[2].text = 'Apellido'
            table.rows[0].cells[3].text = 'Carrera'
            table.rows[0].cells[4].text = 'Maestro'
            table.rows[0].cells[5].text = 'Titulo de titulaciones'
            table.rows[0].cells[6].text = 'Fecha de carga'

            for (numcontrol,nombre,apellido,carrera,maestro,nomarchivo,fecha) in datos2:
                row_cells = table.add_row().cells
                row_cells[0].text = numcontrol
                row_cells[1].text = nombre
                row_cells[2].text = apellido
                row_cells[3].text = carrera
                row_cells[4].text = maestro
                row_cells[5].text = nomarchivo
                row_cells[6].text = fecha

            document.save('REPORTE TITULACIONES {} CARRERA {} RANGO {} A {}.docx'.format(fechaActual,carrera,fechaInicio,fechaFinal))
            path = os.path.realpath('REPORTE TITULACIONES {} CARRERA {} RANGO {} A {}.docx'.format(fechaActual,carrera,fechaInicio,fechaFinal))
            os.startfile(path)

        else:
            ms.showerror('ERROR','Ingresa un rango de fechas o la carrera')
        
    def update_mem(self, numcontrol):
        nombre = self.edit_nombre_mem.get()
        apellido = self.edit_apellido_mem.get()
        cont = self.edit_cont_mem.get()
        carrera = self.edit_carrera_mem.get()
        maestro = self.edit_maestro_mem.get()

        parametros = nombre, apellido, cont, carrera, maestro, numcontrol
        query = 'update memorias set nombre = ? , apellido = ? , pass = ? , carrera = ?, maestro = ? where numcontrol = ?'
        self.run_query(query,parametros)
        ms.showinfo('Actualizacion de datos completada', 'El alumno {} {} {} se actualizo correctamente'.format(numcontrol, nombre, apellido))
        self.obtener_memorias()
    
    def update_tit(self, numcontrol):
        nombre = self.edit_nombre_tit.get()
        apellido = self.edit_apellido_tit.get()
        cont = self.edit_cont_tit.get()
        carrera = self.edit_carrera_tit.get()
        maestro = self.edit_maestro_tit.get()

        parametros = nombre, apellido, cont, carrera, maestro, numcontrol
        query = 'update titulaciones set nombre = ? , apellido = ? , pass = ? , carrera = ?, maestro = ? where numcontrol = ?'
        self.run_query(query,parametros)
        ms.showinfo('Actualizacion de datos completada', 'El alumno {} {} {} se actualizo correctamente'.format(numcontrol, nombre, apellido))
        self.obtener_titulaciones()

### FUNCIONES ALUMNOS   
    def conectar_alumnos_memorias(self):
        u = self.usuario_mem_alumno.get()
        c = self.cont_mem_alumno.get()
        parametros= u,c
        query = 'select * from memorias where numcontrol = ? and pass = ?'
        datos = self.run_query(query,parametros)

        if datos.fetchall():
            self.inicio_memorias()
        else:
            ms.showerror('Error','Usuario Incorrecto')
    
    def conectar_alumnos_titulaciones(self):
        u = self.usuario_tit_alumno.get()
        c = self.cont_tit_alumno.get()
        parametros= u,c
        query = 'select * from titulaciones where numcontrol = ? and pass = ?'
        datos = self.run_query(query,parametros)

        if datos.fetchall():
            self.inicio_titulaciones()
        else:
            ms.showerror('Error','Usuario Incorrecto')

    def reset_password(self):
        ms.showwarning('Olvide mi contraseña','Contacta a tu maestro o administracion para recuperar tu contraseña')

    def seleccion_archivo_mem(self):
        archivo = filedialog.askopenfilename(title = "Seleccionar", initialdir = 'c:/', filetypes = (("PDF", "*.pdf"),("Word", "*.docx")))
        self.archivo_inicio_mem.insert(0,archivo)
    
    def cargar_mem(self):
        usuario = self.numcontrol_inicio_mem.get()
        fecha = date.today()
        nomarchivo = self.nomarchivo_inicio_mem.get()
        archivo = self.archivo_inicio_mem.get()              
        
        os.makedirs('C:\\ITSPR\\Residencias\\{}'.format(usuario), exist_ok=True)
        archivo_ub = 'C:\\ITSPR\\Residencias\\{}'.format(usuario)
        shutil.copy(archivo, archivo_ub)
        nombre = os.listdir(archivo_ub)
        archivo_ub = archivo_ub + '\\' + nombre[0]
        
        parametros = nomarchivo, archivo_ub, fecha, usuario
        query = 'update memorias set nomarchivo = ? , archivo = ? , fecha = ? where numcontrol = ?'
        
        self.run_query(query, parametros)
        ms.showinfo('Archivos cargados correctamente', 'Los datos de {} fueron cargados correctamente'.format(usuario))

    def seleccion_archivo_tit(self):
        archivo = filedialog.askopenfilename(title = "Seleccionar", initialdir = 'c:/', filetypes = (("PDF", "*.pdf"),("Word", "*.docx")))
        self.archivo_inicio_tit.insert(0,archivo)
    
    def cargar_tit(self):
        usuario = self.numcontrol_inicio_tit.get()
        fecha = date.today()
        nomarchivo = self.nomarchivo_inicio_tit.get()
        archivo = self.archivo_inicio_tit.get()              
        
        os.makedirs('C:\\ITSPR\\Titulaciones\\{}'.format(usuario), exist_ok=True)
        archivo_ub = 'C:\\ITSPR\\Titulaciones\\{}'.format(usuario)
        shutil.copy(archivo, archivo_ub)
        nombre = os.listdir(archivo_ub)
        archivo_ub = archivo_ub + '\\' + nombre[0]
        
        parametros = nomarchivo, archivo_ub, fecha, usuario
        query = 'update titulaciones set nomarchivo = ? , archivo = ? , fecha = ? where numcontrol = ?'
        
        self.run_query(query, parametros)
        ms.showinfo('Archivos cargados correctamente', 'Los datos de {} fueron cargados correctamente'.format(usuario))

### MAINLOOP
if __name__ == "__main__":
    ventana = Tk()
    app=Tec(ventana)
    ventana.mainloop()